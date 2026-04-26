"""Built-in template generator — creates default .docx templates with placeholders."""

import os
from sqlalchemy import select
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from config import settings
from models.database import async_session
from models.template import Template


def _add_heading(doc, text: str):
    heading = doc.add_heading(text, level=1)
    for run in heading.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading


def generate_gaoxin_template(output_path: str):
    """Generate built-in 高新技术企业 template."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("国家高新技术企业认定申请书")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("申报企业：{company_name}")
    run.font.size = Pt(14)
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    doc.add_paragraph()

    _add_heading(doc, "一、企业基本情况")
    doc.add_paragraph("{basic}")

    _add_heading(doc, "二、企业研究开发活动情况")
    doc.add_paragraph("{rd}")

    _add_heading(doc, "三、企业知识产权情况")
    doc.add_paragraph("{ip}")

    _add_heading(doc, "四、企业高新技术产品（服务）情况")
    doc.add_paragraph("{product}")

    _add_heading(doc, "五、企业科技人员情况")
    doc.add_paragraph("{staff}")

    _add_heading(doc, "六、企业创新能力评价")
    doc.add_paragraph("{innovation}")

    _add_heading(doc, "七、附件材料清单")
    doc.add_paragraph("{appendix}")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("申报企业（盖章）：{company_name}\n日期：{date}")
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def generate_xiaojuren_template(output_path: str):
    """Generate built-in 专精特新小巨人 template."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('国家级专精特新"小巨人"企业申请书')
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("申报企业：{company_name}")
    run.font.size = Pt(14)
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    doc.add_paragraph()

    _add_heading(doc, "一、企业基本情况")
    doc.add_paragraph("{basic}")

    _add_heading(doc, "二、专业化程度")
    doc.add_paragraph("{specialization}")

    _add_heading(doc, "三、精细化程度")
    doc.add_paragraph("{refinement}")

    _add_heading(doc, "四、特色化程度")
    doc.add_paragraph("{characteristic}")

    _add_heading(doc, "五、创新能力")
    doc.add_paragraph("{innovation}")

    _add_heading(doc, "六、产业链配套能力")
    doc.add_paragraph("{chain}")

    _add_heading(doc, "七、经营管理水平")
    doc.add_paragraph("{management}")

    _add_heading(doc, "八、附件材料清单")
    doc.add_paragraph("{appendix}")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("申报企业（盖章）：{company_name}\n日期：{date}")
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


async def init_builtin_templates():
    """Create built-in .docx templates if they don't already exist in the database."""
    template_dir = os.path.join(settings.data_dir, "templates")
    os.makedirs(template_dir, exist_ok=True)

    builtins = [
        {
            "id": "builtin-gaoxin",
            "name": "高新技术企业认定申请书（内置模板）",
            "project_type": "gaoxin",
            "filename": "高新技术企业认定申请书（模板）.docx",
            "generator": generate_gaoxin_template,
        },
        {
            "id": "builtin-xiaojuren",
            "name": "专精特新小巨人企业申请书（内置模板）",
            "project_type": "xiaojuren",
            "filename": "专精特新小巨人企业申请书（模板）.docx",
            "generator": generate_xiaojuren_template,
        },
    ]

    async with async_session() as db:
        for b in builtins:
            q = select(Template).where(Template.id == b["id"])
            result = await db.execute(q)
            existing = result.scalars().first()
            if existing:
                continue

            file_path = os.path.join(template_dir, b["filename"])
            b["generator"](file_path)

            template = Template(
                id=b["id"],
                name=b["name"],
                project_type=b["project_type"],
                file_path=file_path,
                original_filename=b["filename"],
                is_builtin=True,
                is_active=True,
            )
            db.add(template)

        await db.commit()
