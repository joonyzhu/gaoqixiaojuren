"""Word document exporter — fills a .docx template with generated content.

Placeholders in the template use {key} syntax:
  {company_name} {project_name} {date}
  {basic} {rd} {ip} {product} {staff} {innovation} {appendix}
  {specialization} {refinement} {characteristic} {chain} {management}
"""

import os
import re
from copy import deepcopy
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

PLACEHOLDER_RE = re.compile(r"\{([a-z_]+)\}")

METADATA_DEFAULTS: dict[str, str] = {}


def _current_date():
    return datetime.now().strftime("%Y年%m月%d日")


def _insert_multiline(paragraph, full_text: str):
    """Replace paragraph text and insert subsequent lines as new paragraphs."""
    lines = [l.strip() for l in full_text.split("\n") if l.strip()]
    if not lines:
        return
    # Replace placeholder in current paragraph with first line
    paragraph.text = ""
    run = paragraph.add_run(lines[0])
    _copy_run_font(paragraph.runs[0] if paragraph.runs else run, run)

    # Insert remaining lines as new paragraphs after this one
    parent = paragraph._element.getparent()
    index = list(parent).index(paragraph._element)
    for line in lines[1:]:
        new_p = deepcopy(paragraph._element)
        # Clear existing runs
        for r in new_p.findall(qn('w:r')):
            new_p.remove(r)
        r_elem = new_p.makeelement(qn('w:r'), {})
        t_elem = new_p.makeelement(qn('w:t'), {})
        t_elem.text = line
        t_elem.set(qn('xml:space'), 'preserve')
        r_elem.append(t_elem)
        new_p.append(r_elem)
        index += 1
        parent.insert(index, new_p)


def _copy_run_font(src_run, dst_run):
    """Copy font properties from src to dst."""
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb


def _replace_in_paragraph(para, replacements: dict[str, str]):
    """Replace {placeholders} in a paragraph. Handles placeholders split across
    multiple runs (a common python-docx behavior). For multi-line content,
    inserts extra paragraphs."""
    full_text = "".join(run.text for run in para.runs)
    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    # Process right-to-left so earlier match offsets stay valid after modifications
    for match in reversed(matches):
        key = match.group(1)
        replacement = replacements.get(key)
        if replacement is None:
            continue

        start, end = match.start(), match.end()

        # Find which runs span this placeholder in the CURRENT run texts
        cum = 0
        affected: list[int] = []
        for i, run in enumerate(para.runs):
            run_start = cum
            run_end = cum + len(run.text)
            cum = run_end
            if run_start < end and run_end > start:
                affected.append(i)

        if not affected:
            continue

        # Rebuild the combined text from affected runs (may differ from original
        # full_text if earlier (rightward) replacements already touched these runs)
        combined = "".join(para.runs[i].text for i in affected)

        if "\n" not in replacement:
            combined = combined.replace("{" + key + "}", replacement)
            para.runs[affected[0]].text = combined
            for i in affected[1:]:
                para.runs[i].text = ""
        else:
            lines = [l.strip() for l in replacement.split("\n") if l.strip()]
            if not lines:
                for i in affected:
                    para.runs[i].text = ""
                continue

            combined = combined.replace("{" + key + "}", lines[0])
            para.runs[affected[0]].text = combined
            for i in affected[1:]:
                para.runs[i].text = ""

            if len(lines) > 1:
                parent = para._element.getparent()
                index = list(parent).index(para._element)
                for line in lines[1:]:
                    new_p = deepcopy(para._element)
                    for r in new_p.findall(qn('w:r')):
                        new_p.remove(r)
                    r_elem = new_p.makeelement(qn('w:r'), {})
                    t_elem = new_p.makeelement(qn('w:t'), {})
                    t_elem.text = line
                    t_elem.set(qn('xml:space'), 'preserve')
                    r_elem.append(t_elem)
                    new_p.append(r_elem)
                    index += 1
                    parent.insert(index, new_p)


def _replace_in_table_cell(cell, replacements: dict[str, str]):
    for para in cell.paragraphs:
        _replace_in_paragraph(para, replacements)


def _replace_in_header_footer(section, replacements: dict[str, str]):
    for para in section.header.paragraphs:
        _replace_in_paragraph(para, replacements)
    for para in section.footer.paragraphs:
        _replace_in_paragraph(para, replacements)
    # Also check if header/footer has tables
    for table in section.header.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_table_cell(cell, replacements)
    for table in section.footer.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_table_cell(cell, replacements)


def fill_template(template_path: str, replacements: dict[str, str], output_path: str) -> str:
    """Open a .docx template, replace all {placeholders}, save to output_path."""
    doc = Document(template_path)

    # Replace in body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_table_cell(cell, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        _replace_in_header_footer(section, replacements)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


def export_to_word(
    project_name: str,
    company_name: str,
    project_type: str,
    sections: list[dict],
    output_path: str,
    template_path: str = "",
) -> str:
    """Export all sections as a formatted Word document.

    If template_path is provided, fill the template with content.
    Otherwise use hardcoded default formatting.
    """
    # Build replacement map
    replacements: dict[str, str] = {
        "project_name": project_name,
        "company_name": company_name or "",
        "date": _current_date(),
    }
    for sec in sections:
        replacements[sec.get("id", sec.get("section_id", ""))] = sec.get("content", "")

    # Template-based export
    if template_path and os.path.isfile(template_path):
        return fill_template(template_path, replacements, output_path)

    # Fallback: hardcoded formatting
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    type_label = "国家高新技术企业认定申请书" if project_type == "gaoxin" else '国家级专精特新"小巨人"企业申请书'
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(type_label)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"申报企业：{company_name}\n项目名称：{project_name}")
    run.font.size = Pt(14)
    run.font.name = '楷体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    doc.add_paragraph()

    for sec in sections:
        heading = doc.add_heading(sec.get("title", ""), level=2)
        for run in heading.runs:
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        content = sec.get("content", "")
        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0.74)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(para_text)
                run.font.size = Pt(12)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_paragraph()

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{company_name}\n{_current_date()}")
    run.font.size = Pt(12)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path
