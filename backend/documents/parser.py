"""Multi-format document parser. Supports PDF, DOCX, XLSX, TXT, Markdown, HTML."""

import os
from pathlib import Path


class ParsedDocument:
    def __init__(self, text: str = "", metadata: dict | None = None, pages: list[str] | None = None):
        self.text = text
        self.metadata = metadata or {}
        self.pages = pages or []

    def __repr__(self):
        return f"ParsedDocument(text_len={len(self.text)}, meta={list(self.metadata.keys())})"


def parse_file(file_path: str) -> ParsedDocument:
    """Parse any supported document and return its text content."""
    ext = Path(file_path).suffix.lower()
    metadata = {"file_path": file_path, "file_name": os.path.basename(file_path), "file_type": ext}

    if ext == ".pdf":
        return _parse_pdf(file_path, metadata)
    elif ext in (".docx", ".doc"):
        return _parse_docx(file_path, metadata)
    elif ext in (".xlsx", ".xls"):
        return _parse_xlsx(file_path, metadata)
    elif ext in (".txt", ".text", ".csv"):
        return _parse_text(file_path, metadata)
    elif ext in (".md", ".markdown"):
        return _parse_text(file_path, metadata)
    elif ext in (".html", ".htm"):
        return _parse_html(file_path, metadata)
    else:
        # Try plain text as fallback
        return _parse_text(file_path, metadata)


def _parse_pdf(file_path: str, meta: dict) -> ParsedDocument:
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
        full_text = "\n\n".join(pages)
        meta["page_count"] = len(pages)
        return ParsedDocument(text=full_text, metadata=meta, pages=pages)
    except Exception as e:
        return ParsedDocument(text=f"[PDF解析失败: {e}]", metadata=meta)


def _parse_docx(file_path: str, meta: dict) -> ParsedDocument:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)

        # Also extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            full_text += "\n\n" + "\n".join(rows)

        meta["paragraph_count"] = len(paragraphs)
        meta["table_count"] = len(doc.tables)
        return ParsedDocument(text=full_text, metadata=meta)
    except Exception as e:
        return ParsedDocument(text=f"[DOCX解析失败: {e}]", metadata=meta)


def _parse_xlsx(file_path: str, meta: dict) -> ParsedDocument:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, data_only=True)
        all_text = []
        meta["sheet_count"] = len(wb.sheetnames)
        meta["sheet_names"] = wb.sheetnames

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_text.append(f"=== {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join([str(c) if c is not None else "" for c in row])
                if row_text.strip():
                    all_text.append(row_text)

        return ParsedDocument(text="\n".join(all_text), metadata=meta)
    except Exception as e:
        return ParsedDocument(text=f"[XLSX解析失败: {e}]", metadata=meta)


def _parse_text(file_path: str, meta: dict) -> ParsedDocument:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        meta["char_count"] = len(text)
        return ParsedDocument(text=text, metadata=meta)
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="gbk") as f:
                text = f.read()
            meta["char_count"] = len(text)
            return ParsedDocument(text=text, metadata=meta)
        except Exception as e:
            return ParsedDocument(text=f"[TXT解析失败: {e}]", metadata=meta)


def _parse_html(file_path: str, meta: dict) -> ParsedDocument:
    try:
        from bs4 import BeautifulSoup
        with open(file_path, "r", encoding="utf-8") as f:
            html = f.read()
        soup = BeautifulSoup(html, "html.parser")
        # Remove script and style tags
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        # Clean up excessive blank lines
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return ParsedDocument(text="\n".join(lines), metadata=meta)
    except Exception as e:
        return ParsedDocument(text=f"[HTML解析失败: {e}]", metadata=meta)
